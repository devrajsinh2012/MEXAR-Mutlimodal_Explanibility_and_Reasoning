"""
MEXAR Core Engine - Data Ingestion & Validation Module
Handles parsing and validation of uploaded files (CSV, PDF, DOCX, JSON, TXT).
"""

import os
import json
import logging
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DataValidator:
    """
    Validates and parses uploaded data files for knowledge compilation.
    Supports: CSV, PDF, DOCX, JSON, TXT
    """
    
    # Minimum thresholds for data sufficiency
    MIN_ENTRIES = 20
    MIN_CHARACTERS = 2000
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {'.csv', '.pdf', '.docx', '.json', '.txt'}
    
    def __init__(self):
        """Initialize the data validator."""
        self.parsed_data: List[Dict[str, Any]] = []
        self.validation_results: List[Dict[str, Any]] = []
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a file based on its extension.
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            Dict containing:
                - format: File format (csv, pdf, docx, json, txt)
                - data: Parsed data (list of dicts for structured, None for text)
                - text: Extracted text content
                - entries_count: Number of entries/rows/paragraphs
                - file_name: Original file name
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}. Supported: {self.SUPPORTED_EXTENSIONS}")
        
        result = {
            "format": ext.replace(".", ""),
            "data": None,
            "text": "",
            "entries_count": 0,
            "file_name": path.name
        }
        
        try:
            if ext == '.csv':
                result = self._parse_csv(file_path, result)
            elif ext == '.pdf':
                result = self._parse_pdf(file_path, result)
            elif ext == '.docx':
                result = self._parse_docx(file_path, result)
            elif ext == '.json':
                result = self._parse_json(file_path, result)
            elif ext == '.txt':
                result = self._parse_txt(file_path, result)
                
            logger.info(f"Successfully parsed {path.name}: {result['entries_count']} entries, {len(result['text'])} chars")
            
        except Exception as e:
            logger.error(f"Error parsing {path.name}: {str(e)}")
            result["error"] = str(e)
        
        return result
    
    def _parse_csv(self, file_path: str, result: Dict) -> Dict:
        """Parse CSV file into structured data."""
        df = pd.read_csv(file_path)
        
        # Convert to list of dicts
        data = df.to_dict(orient='records')
        
        # Generate text representation
        text_parts = []
        for i, row in enumerate(data):
            row_text = f"Entry {i+1}: " + ", ".join([f"{k}={v}" for k, v in row.items() if pd.notna(v)])
            text_parts.append(row_text)
        
        result["data"] = data
        result["text"] = "\n".join(text_parts)
        result["entries_count"] = len(data)
        result["columns"] = list(df.columns)
        
        return result
    
    def _parse_pdf(self, file_path: str, result: Dict) -> Dict:
        """Parse PDF file and extract text."""
        reader = PdfReader(file_path)
        
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"Page {i+1}:\n{page_text}")
        
        full_text = "\n\n".join(text_parts)
        
        # Count paragraphs as entries
        paragraphs = [p.strip() for p in full_text.split('\n\n') if p.strip()]
        
        result["text"] = full_text
        result["entries_count"] = len(paragraphs)
        result["page_count"] = len(reader.pages)
        
        return result
    
    def _parse_docx(self, file_path: str, result: Dict) -> Dict:
        """Parse DOCX file and extract text."""
        doc = Document(file_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Also extract tables
        table_data = []
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):
                    table_data.append(row_data)
        
        result["text"] = "\n\n".join(paragraphs)
        result["entries_count"] = len(paragraphs) + len(table_data)
        result["table_data"] = table_data
        
        return result
    
    def _parse_json(self, file_path: str, result: Dict) -> Dict:
        """Parse JSON file into structured data."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Handle different JSON structures
        if isinstance(data, list):
            entries = data
        elif isinstance(data, dict):
            # If it's a dict with a main data key, extract it
            for key in ['data', 'items', 'records', 'entries']:
                if key in data and isinstance(data[key], list):
                    entries = data[key]
                    break
            else:
                # Wrap single object in list
                entries = [data]
        else:
            entries = [{"value": data}]
        
        # Generate text representation
        text_parts = []
        for i, entry in enumerate(entries):
            if isinstance(entry, dict):
                entry_text = f"Entry {i+1}: " + json.dumps(entry, ensure_ascii=False)
            else:
                entry_text = f"Entry {i+1}: {entry}"
            text_parts.append(entry_text)
        
        result["data"] = entries
        result["text"] = "\n".join(text_parts)
        result["entries_count"] = len(entries)
        
        return result
    
    def _parse_txt(self, file_path: str, result: Dict) -> Dict:
        """Parse TXT file as plain text."""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Count lines as entries
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        result["text"] = text
        result["entries_count"] = len(lines)
        
        return result
    
    def validate_sufficiency(self, parsed_data: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Check if the combined data meets minimum requirements.
        
        Args:
            parsed_data: List of parsed file results
            
        Returns:
            Dict containing:
                - sufficient: Boolean indicating if data is sufficient
                - issues: List of issues found
                - warnings: List of warnings
                - stats: Statistics about the data
        """
        total_entries = sum(p.get("entries_count", 0) for p in parsed_data)
        total_chars = sum(len(p.get("text", "")) for p in parsed_data)
        
        issues = []
        warnings = []
        
        # Check minimum thresholds
        entries_ok = total_entries >= self.MIN_ENTRIES
        chars_ok = total_chars >= self.MIN_CHARACTERS
        
        if not entries_ok and not chars_ok:
            issues.append(
                f"Insufficient data: Found {total_entries} entries and {total_chars} characters. "
                f"Need at least {self.MIN_ENTRIES} entries OR {self.MIN_CHARACTERS} characters."
            )
        
        # Check for empty files
        empty_files = [p["file_name"] for p in parsed_data if p.get("entries_count", 0) == 0]
        if empty_files:
            issues.append(f"Empty or unreadable files: {', '.join(empty_files)}")
        
        # Check for parsing errors
        error_files = [p["file_name"] for p in parsed_data if "error" in p]
        if error_files:
            issues.append(f"Files with parsing errors: {', '.join(error_files)}")
        
        # Add warnings for low-quality data
        if total_entries < self.MIN_ENTRIES * 2:
            warnings.append(
                f"Consider adding more entries for better knowledge coverage. "
                f"Current: {total_entries}, Recommended: {self.MIN_ENTRIES * 2}+"
            )
        
        # Calculate structure score (how well-structured the data is)
        structured_count = sum(1 for p in parsed_data if p.get("data") is not None)
        structure_score = structured_count / len(parsed_data) if parsed_data else 0
        
        if structure_score < 0.5:
            warnings.append(
                "Most files are unstructured (PDF/TXT). "
                "Structured data (CSV/JSON) provides better knowledge extraction."
            )
        
        # Compile statistics
        stats = {
            "total_files": len(parsed_data),
            "total_entries": total_entries,
            "total_characters": total_chars,
            "structure_score": round(structure_score, 2),
            "file_breakdown": [
                {
                    "name": p["file_name"],
                    "format": p["format"],
                    "entries": p.get("entries_count", 0),
                    "characters": len(p.get("text", ""))
                }
                for p in parsed_data
            ]
        }
        
        return {
            "sufficient": len(issues) == 0,
            "issues": issues,
            "warnings": warnings,
            "stats": stats
        }
    
    def provide_feedback(self, validation_result: Dict[str, Any]) -> str:
        """
        Generate user-friendly feedback message.
        
        Args:
            validation_result: Result from validate_sufficiency
            
        Returns:
            Formatted feedback message
        """
        stats = validation_result["stats"]
        
        if validation_result["sufficient"]:
            # Success message
            feedback = f"""✅ **Data Validation Passed!**

📊 **Statistics:**
- Total Files: {stats['total_files']}
- Total Entries: {stats['total_entries']}
- Total Characters: {stats['total_characters']:,}
- Structure Score: {stats['structure_score']*100:.0f}%

"""
            # Add file breakdown
            feedback += "📁 **File Breakdown:**\n"
            for f in stats["file_breakdown"]:
                feedback += f"- {f['name']} ({f['format'].upper()}): {f['entries']} entries\n"
            
            # Add warnings if any
            if validation_result["warnings"]:
                feedback += "\n⚠️ **Suggestions:**\n"
                for warning in validation_result["warnings"]:
                    feedback += f"- {warning}\n"
        
        else:
            # Failure message
            feedback = f"""❌ **Data Validation Failed**

🔍 **Issues Found:**
"""
            for issue in validation_result["issues"]:
                feedback += f"- {issue}\n"
            
            feedback += f"""
📊 **Current Statistics:**
- Total Entries: {stats['total_entries']} (minimum: {self.MIN_ENTRIES})
- Total Characters: {stats['total_characters']:,} (minimum: {self.MIN_CHARACTERS:,})

💡 **How to Fix:**
1. Add more data files (CSV, PDF, DOCX, JSON, or TXT)
2. Ensure files contain meaningful content
3. For best results, use structured formats like CSV or JSON
"""
        
        return feedback
    
    def parse_and_validate(self, file_paths: List[str]) -> Tuple[List[Dict], Dict, str]:
        """
        Convenience method to parse all files and validate in one call.
        
        Args:
            file_paths: List of file paths to process
            
        Returns:
            Tuple of (parsed_data, validation_result, feedback_message)
        """
        parsed_data = []
        for path in file_paths:
            result = self.parse_file(path)
            parsed_data.append(result)
        
        validation = self.validate_sufficiency(parsed_data)
        feedback = self.provide_feedback(validation)
        
        return parsed_data, validation, feedback


# Factory function for easy instantiation
def create_validator() -> DataValidator:
    """Create a new DataValidator instance."""
    return DataValidator()
